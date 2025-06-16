import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def create_job_report_docx(csv_file_path: str, output_docx_path: str):
    """
    Reads job data from a CSV file and generates a structured .docx report.
    """
    try:
        jobs_df = pd.read_csv(csv_file_path)
        logging.info(f"Successfully loaded {len(jobs_df)} jobs from {csv_file_path}")
    except FileNotFoundError:
        logging.error(f"Error: The file {csv_file_path} was not found.")
        return
    except pd.errors.EmptyDataError:
        logging.error(f"Error: The file {csv_file_path} is empty.")
        return
    except Exception as e:
        logging.error(f"An error occurred while reading {csv_file_path}: {e}")
        return

    doc = Document()

    # Add a title to the document
    doc.add_heading('Job Application Report', level=0)
    doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Jobs in Database: {len(jobs_df)}")
    doc.add_page_break()

    # Filter for jobs that are either 'applied' or have a relevance score >= 70 
    # (these are the ones the user is most interested in for a report)
    # Ensure 'relevance_score' is numeric for comparison
    jobs_df['relevance_score'] = pd.to_numeric(jobs_df['relevance_score'], errors='coerce').fillna(0)
    
    reportable_jobs = jobs_df[
        (jobs_df['status'] == 'applied') | 
        ((jobs_df['status'] == 'new') & (jobs_df['relevance_score'] >= 70)) |
        (jobs_df['status'] == 'reviewed - decided') # Include jobs reviewed but not applied for
    ].copy() # Create a copy to avoid SettingWithCopyWarning

    logging.info(f"Found {len(reportable_jobs)} jobs to include in the report.")

    if reportable_jobs.empty:
        doc.add_paragraph("No jobs met the criteria for reporting (status 'applied' or score >= 70 and status 'new' or 'reviewed - decided').")
    else:
        # Sort by relevance score (descending) then by discovery date (descending)
        reportable_jobs.sort_values(by=['relevance_score', 'discovery_date'], ascending=[False, False], inplace=True)

        for index, job in reportable_jobs.iterrows():
            try:
                title = job.get('title', 'N/A')
                company = job.get('company', 'N/A')
                
                heading_text = f"{title} at {company}"
                doc.add_heading(heading_text, level=1)

                doc.add_paragraph().add_run(f"Job ID: ").bold = True
                doc.add_paragraph(str(job.get('job_id', 'N/A')))
                
                doc.add_paragraph().add_run(f"Status: ").bold = True
                doc.add_paragraph(str(job.get('status', 'N/A')))

                doc.add_paragraph().add_run(f"Location: ").bold = True
                doc.add_paragraph(str(job.get('location', 'N/A')))

                doc.add_paragraph().add_run(f"Job Type: ").bold = True
                doc.add_paragraph(str(job.get('job_type', 'N/A')))

                doc.add_paragraph().add_run(f"Source: ").bold = True
                doc.add_paragraph(str(job.get('source', 'N/A')))

                doc.add_paragraph().add_run(f"Discovery Date: ").bold = True
                doc.add_paragraph(str(job.get('discovery_date', 'N/A')))
                
                application_link = job.get('application_link', job.get('link', 'N/A'))
                doc.add_paragraph().add_run(f"Application Link: ").bold = True
                if pd.notna(application_link):
                    # Add hyperlink (basic approach, advanced might require more complex XML)
                    p = doc.add_paragraph()
                    p.add_run(application_link) #.hyperlink = application_link # Hyperlink property not directly settable on run
                else:
                    doc.add_paragraph('N/A')


                doc.add_paragraph().add_run(f"Relevance Score (Gemini): ").bold = True
                doc.add_paragraph(f"{job.get('relevance_score', 'N/A')}%")

                # Handle Gemini Analysis (assuming it's stored as a JSON string)
                doc.add_paragraph().add_run("Gemini Analysis:").bold = True
                gemini_analysis_str = job.get('gemini_analysis')
                if pd.notna(gemini_analysis_str):
                    try:
                        analysis_data = json.loads(gemini_analysis_str)
                        if isinstance(analysis_data, dict):
                            doc.add_paragraph(f"  Recommendation: {analysis_data.get('overall_recommendation', 'N/A')}")
                            doc.add_paragraph(f"  Summary: {analysis_data.get('summary_for_candidate', 'N/A')}")
                            
                            strengths = analysis_data.get('key_strengths_alignment', [])
                            if strengths:
                                doc.add_paragraph("  Key Strengths Alignment:")
                                for item in strengths: doc.add_paragraph(f"    - {item}")
                            
                            gaps = analysis_data.get('potential_gaps_or_areas_for_development', [])
                            if gaps:
                                doc.add_paragraph("  Potential Gaps/Areas for Development:")
                                for item in gaps: doc.add_paragraph(f"    - {item}")
                            
                            strategy = analysis_data.get('application_strategy_suggestions', [])
                            if strategy:
                                doc.add_paragraph("  Application Strategy Suggestions:")
                                for item in strategy: doc.add_paragraph(f"    - {item}")
                        else: # If not a dict after loading (e.g. just a string error message from Gemini)
                            doc.add_paragraph(f"  {str(analysis_data)}")
                    except json.JSONDecodeError:
                        doc.add_paragraph(f"  Could not parse Gemini analysis: {gemini_analysis_str}")
                    except Exception as e_parse:
                        doc.add_paragraph(f"  Error displaying Gemini analysis: {str(e_parse)}")
                else:
                    doc.add_paragraph("  N/A")
                
                doc.add_paragraph("-" * 80) # Separator
                # doc.add_page_break() # Optional: new page for each job

            except Exception as e_job:
                logging.error(f"Error processing job {job.get('job_id', 'Unknown ID')}: {e_job}")
                doc.add_paragraph(f"Could not fully process job ID: {job.get('job_id', 'N/A')}. Error: {e_job}")
                continue
                
    try:
        doc.save(output_docx_path)
        logging.info(f"Job report successfully saved to {output_docx_path}")
    except Exception as e:
        logging.error(f"Error saving Word document: {e}")

if __name__ == '__main__':
    csv_file = Path('jobs_database.csv')
    docx_output_file = Path('job_application_report.docx')
    create_job_report_docx(str(csv_file), str(docx_output_file)) 